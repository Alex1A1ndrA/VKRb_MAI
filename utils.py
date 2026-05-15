import re
import random
import pandas as pd
from datetime import datetime, timedelta
import pymorphy3
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import streamlit as st

from constants import DASH

morph = pymorphy3.MorphAnalyzer()

def extract_group(text: str) -> str | None:
    match = re.search(r'\d{3}', text)
    return match.group(0) if match else None

def value_or_dash(val):
    if pd.isna(val) or str(val).strip() == "":
        return DASH
    return str(val).strip()

def safe_date_conversion(date_str, dayfirst=True):
    if pd.isna(date_str):
        return pd.NaT
    try:
        return pd.to_datetime(date_str, dayfirst=dayfirst, errors='coerce')
    except:
        return pd.NaT

def get_parsed_word(word: str, gender: str = None):
    parsed_list = morph.parse(word)
    if not parsed_list:
        return None
    if gender is None:
        return parsed_list[0]
    for p in parsed_list:
        if p.tag.gender == gender:
            return p
    return parsed_list[0]

def inflect_word(word: str, case: str, gender: str = None) -> str:
    if not word:
        return word
    parsed = get_parsed_word(word, gender)
    if parsed is None:
        return word
    tags = {case}
    if gender:
        tags.add(gender)
    inflected = parsed.inflect(tags)
    if inflected:
        return inflected.word
    return word

def detect_gender_by_patronymic(fio: str):
    parts = fio.strip().split()
    if len(parts) != 3:
        return None
    patr = parts[2].lower()
    if patr.endswith(('ич', 'вич')):
        return 'masc'
    if patr.endswith(('на', 'овна', 'евна')):
        return 'femn'
    return None

def manual_surname_rules(surname: str, case: str, gender: str) -> str | None:
    s = surname.strip()
    s_low = s.lower().replace('ё', 'е')
    if case == 'nomn':
        return s
    if gender == 'femn' and s_low.endswith('ская'):
        core = s[:-2]
        if case in ('gent', 'datv', 'ablt', 'loct'):
            return core + 'ой'
        if case == 'accs':
            return core + 'ую'
    if gender == 'femn' and s_low.endswith(('ина', 'ова', 'ева')):
        core = s[:-1]
        if case == 'gent':
            return core + 'ой'
        if case == 'datv':
            return core + 'ой'
        if case == 'accs':
            return core + 'у'
        if case in ('ablt', 'loct'):
            return core + 'ой'
    if gender == 'femn' and s_low.endswith('ая'):
        core = s[:-2]
        if case in ('gent', 'datv', 'ablt', 'loct'):
            return core + 'ой'
        if case == 'accs':
            return core + 'у'
    if gender == 'masc':
        if s_low.endswith(('ов', 'ев', 'ёв', 'ин', 'ын')):
            if case in ('gent', 'accs'):
                return s + 'а'
            if case == 'datv':
                return s + 'у'
            if case == 'ablt':
                return s + 'ым'
            if case == 'loct':
                return s + 'е'
        if s_low.endswith(('ский', 'ской')):
            core = s[:-4]
            if case == 'gent':
                return core + 'ского'
            if case == 'datv':
                return core + 'скому'
            if case == 'ablt':
                return core + 'ским'
            if case == 'loct':
                return core + 'ском'
        if s_low.endswith('ой'):
            core = s[:-2]
            if case == 'gent':
                return core + 'ого'
            if case == 'datv':
                return core + 'ому'
            if case == 'ablt':
                return core + 'ым'
            if case == 'loct':
                return core + 'ом'
        if s_low.endswith('а'):
            core = s[:-1]
            if case == 'gent':
                return core + 'ы'
            if case == 'datv':
                return core + 'е'
            if case == 'ablt':
                return core + 'ой'
            if case == 'loct':
                return core + 'е'
    return None

def fio_in_case(fio: str, case: str, gender: str = None) -> str:
    parts = fio.strip().split()
    if len(parts) != 3:
        return fio
    fam, name, patr = parts
    if not gender:
        gender = detect_gender_by_patronymic(fio)
    if gender is None:
        return fio
    name_res = inflect_word(name, case, gender).capitalize()
    patr_res = inflect_word(patr, case, gender).capitalize()
    fam_manual = manual_surname_rules(fam, case, gender)
    if fam_manual is not None:
        fam_res = fam_manual.capitalize()
    else:
        fam_res = inflect_word(fam, case, gender).capitalize()
    return f"{fam_res} {name_res} {patr_res}"
    return f"{fam_res} {name_res} {patr_res}"

def short_fio(fio):
    if pd.isna(fio) or str(fio).strip() == "":
        return ""
    parts = str(fio).strip().split()
    if len(parts) != 3:
        return fio
    return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."

def role_fio_case(fulltext, case='gent'):
    if pd.isna(fulltext) or str(fulltext).strip() == "":
        return ""
    text = str(fulltext).split(',')[0].strip()
    parts = text.strip().split()
    if len(parts) < 3:
        return text
    gender = detect_gender_by_patronymic(text)
    fio_case = fio_in_case(text, case, gender)
    parts = fio_case.split()
    initials = f"{parts[1][0].upper()}.{parts[2][0].upper()}."
    return f"{parts[0]} {initials}"

def clean_text(text):
    if pd.isna(text):
        return ""
    return ' '.join(str(text).split())

def parse_questions(raw, allowed):
    if pd.isna(raw):
        raw = ''
    lines = str(raw).split('\n')
    result = []
    commiss_clean = [p for p in allowed if isinstance(p, str) and p.strip() != '']
    for line in lines:
        if '-' not in line:
            continue
        try:
            asker, text_q = [s.strip() for s in line.split('-', 1)]
            found = False
            for person in commiss_clean:
                if person.split()[0] in asker:
                    result.append(f"{short_fio(person)} - {text_q}")
                    found = True
                    break
            if not found and commiss_clean:
                rand_person = random.choice(commiss_clean)
                result.append(f"{short_fio(rand_person)} - {text_q}")
            if len(result) == 4:
                break
        except:
            continue
    while len(result) < 4:
        result.append('')
    return result

def remove_empty_paragraphs_around_tables(docx_path):
    doc = Document(docx_path)
    new_elements = []
    prev_is_table = False
    sectPr = None

    for element in doc.element.body:
        tag = element.tag
        if 'tbl' in tag:
            while new_elements and new_elements[-1].tag.endswith('p'):
                last_para = new_elements[-1]
                text = ''.join(node.text or '' for node in last_para.iter())
                if text.strip() == '':
                    new_elements.pop()
                else:
                    break
            new_elements.append(element)
            prev_is_table = True
        elif 'p' in tag:
            text = ''.join(node.text or '' for node in element.iter())
            if text.strip() == '' and prev_is_table:
                continue
            new_elements.append(element)
            prev_is_table = False
        elif 'sectPr' in tag:
            sectPr = element
        else:
            new_elements.append(element)
            prev_is_table = False

    if sectPr is not None:
        new_elements.append(sectPr)

    body = doc.element.body
    for _ in range(len(body)):
        body.remove(body[0])
    for el in new_elements:
        body.append(el)
    doc.save(docx_path)

def insert_empty_paragraphs_after_text(docx_path, search_text, count=2):
    doc = Document(docx_path)
    for para in doc.paragraphs:
        if search_text in para.text:
            para_element = para._element
            parent = para_element.getparent()
            for _ in range(count):
                new_p = Document().add_paragraph()._element
                parent.insert(parent.index(para_element) + 1, new_p)
            doc.save(docx_path)
            return

def set_landscape(doc):
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

def ensure_section(doc, width, height):
    if not doc.sections:
        sectPr = OxmlElement('w:sectPr')
        pgSz = OxmlElement('w:pgSz')
        pgSz.set(qn('w:w'), str(width.twips))
        pgSz.set(qn('w:h'), str(height.twips))
        pgSz.set(qn('w:orient'), 'landscape')
        sectPr.append(pgSz)
        doc.element.body.append(sectPr)
    else:
        sect = doc.sections[0]
        if sect.page_width.inches < 1 or sect.page_height.inches < 1:
            sectPr = sect._sectPr
            old_pgSz = sectPr.find(qn('w:pgSz'))
            if old_pgSz is not None:
                sectPr.remove(old_pgSz)
            pgSz = OxmlElement('w:pgSz')
            pgSz.set(qn('w:w'), str(width.twips))
            pgSz.set(qn('w:h'), str(height.twips))
            pgSz.set(qn('w:orient'), 'landscape')
            sectPr.append(pgSz)

def parse_defense_dates(date_str):
    if pd.isna(date_str):
        return []
    main_part = str(date_str).split(',')[0].strip()
    main_part = main_part.replace('\\', ',').replace('/', ',')
    dates = [d.strip() for d in main_part.split(',') if d.strip()]
    return dates

def parse_predefense_dates(date_str):
    if pd.isna(date_str):
        return [], ''
    time_match = re.search(r'<br>(\d{1,2}:\d{2})', date_str)
    time_str = time_match.group(1) if time_match else ''
    main_part = date_str.split('<br>')[0].strip()
    dates = re.findall(r'\d{2}\.\d{2}\.\d{2}', main_part)
    if not dates:
        return [], time_str
    if '-' in main_part and len(dates) >= 2:
        return dates[:2], time_str
    else:
        return [dates[0]], time_str

def split_students_by_dates(students_df, dates_list):
    n = len(dates_list)
    if n == 0:
        return []
    total = len(students_df)
    if total == 0:
        return [[] for _ in range(n)]
    indices = list(range(total))
    if n == 1:
        return [indices]
    base = total // n
    remainder = total % n
    sizes = [base + 1 if i < remainder else base for i in range(n)]
    result = []
    start = 0
    for size in sizes:
        result.append(indices[start:start+size])
        start += size
    return result